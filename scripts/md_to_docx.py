#!/usr/bin/env python
"""Convert the post to .docx with headings, tables, lists and figures intact.

Written against the document's own structure rather than by round-tripping HTML,
so headings stay headings, tables stay tables, and the figures are placed at
their referenced positions rather than dropped.
"""
import argparse
import io
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


INLINE = re.compile(
    r"(\[[^\]]+\]\([^)\s]+\)|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
)
LINK = re.compile(r"^\[([^\]]+)\]\(([^)\s]+)\)$")


def _style(run, bold, italic, code):
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)


def add_hyperlink(par, url, text, bold, italic):
    """Add an external hyperlink, keeping any emphasis inside its text.

    python-docx has no hyperlink API, so the relationship and the w:hyperlink
    element are built directly. The runs are created on the paragraph first and
    then moved inside the element, which keeps them in document order because
    the element is appended at the point the link is reached.
    """
    r_id = par.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyper = OxmlElement("w:hyperlink")
    hyper.set(qn("r:id"), r_id)

    first = len(par.runs)
    add_runs(par, text, bold=bold, italic=italic)
    for run in par.runs[first:]:
        run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
        run.font.underline = True
        hyper.append(run._r)
    par._p.append(hyper)


def add_runs(par, text, bold=False, italic=False, code=False):
    """Render links, bold, italic and code spans as runs.

    Each branch recurses on the stripped text rather than emitting it directly,
    so nesting works: a link inside bold, or an italic title inside a link.
    """
    for piece in INLINE.split(text):
        if not piece:
            continue
        m = LINK.match(piece)
        if m:
            add_hyperlink(par, m.group(2), m.group(1), bold, italic)
        elif piece.startswith("**") and piece.endswith("**"):
            add_runs(par, piece[2:-2], True, italic, code)
        elif piece.startswith("*") and piece.endswith("*"):
            add_runs(par, piece[1:-1], bold, True, code)
        elif piece.startswith("`") and piece.endswith("`"):
            _style(par.add_run(piece[1:-1]), bold, italic, True)
        else:
            _style(par.add_run(piece), bold, italic, code)


def cells(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--src", required=True)
    p.add_argument("--out", required=True)
    a = p.parse_args()

    base = os.path.dirname(os.path.abspath(a.src))
    lines = io.open(a.src, encoding="utf-8").read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(10)

    n_tab = n_img = 0
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        if s in ("---", "***", "___"):
            i += 1
            continue

        # image
        m = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", s)
        if m:
            alt, rel = m.group(1), m.group(2)
            path = os.path.join(base, rel)
            if os.path.isfile(path):
                doc.add_picture(path, width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(alt)
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                n_img += 1
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            doc.add_heading(m.group(2).replace("**", ""), level=len(m.group(1)))
            i += 1
            continue

        # table: header row, separator, body
        if s.startswith("|") and i + 1 < len(lines) and \
           re.fullmatch(r"\|(\s*:?-+:?\s*\|)+", lines[i + 1].strip()):
            header = cells(line)
            i += 2
            body = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                body.append(cells(lines[i]))
                i += 1
            tab = doc.add_table(rows=1, cols=len(header))
            tab.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                c = tab.rows[0].cells[j]
                c.text = ""
                add_runs(c.paragraphs[0], h)
                for run in c.paragraphs[0].runs:
                    run.bold = True
            for row in body:
                r = tab.add_row().cells
                for j, val in enumerate(row[:len(header)]):
                    r[j].text = ""
                    add_runs(r[j].paragraphs[0], val)
            for row in tab.rows:
                for c in row.cells:
                    for par in c.paragraphs:
                        par.paragraph_format.space_after = Pt(2)
                        for run in par.runs:
                            run.font.size = Pt(9)
            doc.add_paragraph()
            n_tab += 1
            continue

        # list
        if re.match(r"^[-*+]\s+", s):
            par = doc.add_paragraph(style="List Bullet")
            add_runs(par, re.sub(r"^[-*+]\s+", "", s))
            i += 1
            continue

        # blockquote
        if s.startswith(">"):
            par = doc.add_paragraph(style="Intense Quote")
            add_runs(par, s.lstrip("> ").strip())
            i += 1
            continue

        # subtitle: the italic line right under the title
        if s.startswith("*") and s.endswith("*") and i < 4:
            par = doc.add_paragraph()
            r = par.add_run(s.strip("*"))
            r.italic = True
            r.font.size = Pt(11.5)
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # Body paragraph: markdown joins consecutive non-blank lines, so
        # gather them before rendering. Stop at a blank line or at anything
        # that starts a block of its own -- otherwise a wrapped paragraph
        # becomes several stacked ones and emphasis spanning a line break
        # loses its markers.
        buf = [s]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt or nxt in ("---", "***", "___"):
                break
            if re.match(r"^(#{1,4})\s+", nxt) or re.match(r"^[-*+]\s+", nxt):
                break
            if nxt.startswith(">") or nxt.startswith("|") or nxt.startswith("```"):
                break
            if re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", nxt):
                break
            buf.append(nxt)
            j += 1
        par = doc.add_paragraph()
        add_runs(par, " ".join(buf))
        i = j

    doc.save(a.out)
    print("tables: %d, images: %d" % (n_tab, n_img))
    print("wrote %s (%.2f MB)" % (a.out, os.path.getsize(a.out) / 1e6))


if __name__ == "__main__":
    main()
